"""
Create Administrative Report 的功能在此文件中一并实现

begin 和 end 都是 八位 如 20210101
这个文件里的函数对日期没有任何检错机制 请前端注意！！

return_total_hours_worked(employee_ID, begin, end):
返回某员工的总工作时间

return_total_hours_for_a_project(employee_ID, charge_number, begin, end)
返回某员工为某charge number的工作时间
没有任何检错机制 因为输入值在前端会进行筛选

create_ONE_report(employee_ID, charge_number, begin, end, workhour)
创建一条员工报告 如果要多条 员工报告 将返回值相加即可

create_doc(user, location, title, info_tuple)
创建word
注意 这里的user是操作者的employee ID
location 是文件保存路径（文件名已经包含在其中 结尾是.docx）
title 是文档的标题（注意 不是文件名 不是文件名 不是文件名 是打开word之后看到的文档的题目）
info_tuple 是员工信息 可以包含多条

"""
from datetime import datetime
import pymysql
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT


def cut(num, c):
    str_num = str(num)
    return float(str_num[:str_num.index('.') + 1 + c])


def isTrueDate(a):
    # 查询输入日期是否有效
    # 把年月日剥离出来
    year = int(a[0:4])
    month = int(a[4:6])
    day = int(a[6:])

    print(year)
    print(month)
    print(day)

    if year <= 0 or year > 9999:
        return False
    elif month <= 0 or month > 12:
        return False

    if month == 1 or month == 3 or month == 5 or month == 7 or month == 8 or month == 10 or month == 12:
        if day > 31 or day < 1:
            return False
        else:
            return True

        # 四六九冬
    elif month == 4 or month == 6 or month == 9 or month == 11:
        if day > 30 or day < 1:
            return False
        else:
            return True
        # 平年二月二十八
    else:
        if year % 4 == 0 and year % 400 == 0:
            print('leap')
            if day > 29 or day < 1:
                return False
            else:
                return True

        elif year % 4 == 0 and year % 100 != 0:
            print('leap')
            if day > 29 or day < 1:
                return False
            else:
                return True
        else:
            print('common')
            if day > 28 or day < 1:
                return False
            else:
                return True


def return_total_hours_worked(employee_ID, begin_year, begin_month, begin_day, end_year, end_month, end_day):

    con = pymysql.Connect(
        host='XX.XXX.XXX.XX',
        port=XXXX,
        user='XXX',
        passwd='XXXXXXXXXX',
        db='XXX'
    )
    cur = con.cursor()
    sql = "select * from employee_info where employee_type <> 'PA' and  "+'employee_ID = ' + "'" + employee_ID + "'"
    print(sql)
    cur.execute(sql)
    value = cur.fetchall()
    cur.close()
    con.close()

    if value == ():
        print('employee_ID不存在')
        return '员工不存在！'

    if employee_ID.isspace() or begin_year.isspace() or begin_month.isspace() \
            or begin_day.isspace() or end_year.isspace() \
            or end_month.isspace() or end_day.isspace() == True:
        print('无效输入(1)空格')
        return '输入内容不能全为空格！'
    elif len(employee_ID) == 0 or len(begin_year) == 0 or len(begin_month) == 0 \
            or len(begin_day) == 0 or len(end_year) == 0\
            or len(end_month) == 0 or len(end_day) == 0:
        print('无效输入(2)空值')
        return '输入内容不能为空！'

    elif len(begin_year) > 50:
        print('begin_year太长')
        return '开始年份过长！'
    elif len(begin_month) > 50:
        print('begin_month太长')
        return '开始月份过长！'
    elif len(begin_day) > 50:
        print('begin_day太长')
        return '开始日过长！'
    elif len(end_year) > 50:
        print('end_year太长')
        return '结束年份过长！'
    elif len(end_month) > 50:
        print('end_month太长')
        return '结束月份过长！'
    elif len(end_day) > 50:
        print('end_day太长')
        return '结束日过长！'

    if (begin_year.isdigit() and begin_month.isdigit() and begin_day.isdigit()) == False:
        print("begin date is not digit!")
        return "无效的开始日期！"

    elif '.' in begin_year or '.' in begin_month or '.' in begin_day == True:
        print("no dot!")
        return "无效的开始日期！"

    if (end_year.isdigit() and end_month.isdigit() and end_day.isdigit()) == False:
        print("end date is not digit!")
        return "无效的结束日期！"

    elif '.' in end_year or '.' in end_month or '.' in end_day == True:
        print("no dot!")
        return "无效的结束日期！"

    elif int(begin_year) <= 0 or int(begin_year) > 9999:
        print("begin_year <= 0 or > 9999")
        return "无效的开始年份！"
    elif int(begin_month) <= 0 or int(begin_month) > 12:
        print("begin_month <= 0 or > 12")
        return "无效的开始月份！"
    elif int(begin_day) <= 0 or int(begin_day) > 31:
        print("begin_day <= 0 or > 31")
        return "无效的开始日！"
    elif int(end_year) <= 0 or int(end_year) > 9999:
        print("end_year <= 0 or > 9999")
        return "无效的结束年份！"
    elif int(end_month) <= 0 or int(end_month) > 12:
        print("end_month <= 0 or > 12")
        return "无效的结束月份！"
    elif int(end_day) <= 0 or int(end_day) > 31:
        print("end_day <= 0 or > 31")
        return "无效的结束日！"

    begin = str(int(begin_year)).zfill(4) + str(int(begin_month)).zfill(2) + str(int(begin_day)).zfill(2)
    end = str(int(end_year)).zfill(4) + str(int(end_month)).zfill(2) + str(int(end_day)).zfill(2)

    if isTrueDate(begin) == False:
        print("invalid begin date")
        return "无效的开始日期！"
    if isTrueDate(end) == False:
        print("invalid end date")
        return "无效的结束日期！"

    if begin > end:
        print("begin不能大于end")
        return "无效的日期！"

    begindate = employee_ID + begin
    enddate = employee_ID + end

    con = pymysql.Connect(
        host='XX.XXX.XXX.XX',
        port=XXXX,
        user='XXX',
        passwd='XXXXXXXXXX',
        db='XXX'
    )
    cur = con.cursor()
    sql = 'select work_hour from timecards ' +\
          'where ' + 'employee_ID_date <= ' + "'" + enddate + "' " +\
          "and employee_ID_date >= '" + begindate + "'"
    print(sql)
    cur.execute(sql)
    value = cur.fetchall()
    print(value)
    cur.close()
    con.close()

    total = 0.0
    for i in value:
        # print(i)
        total += i[0]
    return total


def return_total_hours_for_a_project(employee_ID, charge_number, begin_year, begin_month, begin_day, end_year, end_month, end_day):

    con = pymysql.Connect(
        host='XX.XXX.XXX.XX',
        port=XXXX,
        user='XXX',
        passwd='XXXXXXXXXX',
        db='XXX'
    )
    cur = con.cursor()
    sql = "select * from employee_info where employee_type <> 'PA' and  "+'employee_ID = ' + "'" + employee_ID + "'"
    print(sql)
    cur.execute(sql)
    value = cur.fetchall()
    cur.close()
    con.close()

    if value == ():
        print('employee_ID不存在')
        return '员工不存在！'

    if employee_ID.isspace() or begin_year.isspace() or begin_month.isspace() \
            or begin_day.isspace() or end_year.isspace() \
            or end_month.isspace() or end_day.isspace() == True:
        print('无效输入(1)空格')
        return '输入内容不能全为空格！'
    elif len(employee_ID) == 0 or len(begin_year) == 0 or len(begin_month) == 0 \
            or len(begin_day) == 0 or len(end_year) == 0 \
            or len(end_month) == 0 or len(end_day) == 0:
        print('无效输入(2)空值')
        return '输入内容不能为空！'

    elif len(begin_year) > 50:
        print('begin_year太长')
        return '开始年份过长！'
    elif len(begin_month) > 50:
        print('begin_month太长')
        return '开始月份过长！'
    elif len(begin_day) > 50:
        print('begin_day太长')
        return '开始日过长！'
    elif len(end_year) > 50:
        print('end_year太长')
        return '结束年份过长！'
    elif len(end_month) > 50:
        print('end_month太长')
        return '结束月份过长！'
    elif len(end_day) > 50:
        print('end_day太长')
        return '结束日过长！'

    if (begin_year.isdigit() and begin_month.isdigit() and begin_day.isdigit()) == False:
        print("begin date is not digit!")
        return "无效的开始日期！"

    elif '.' in begin_year or '.' in begin_month or '.' in begin_day == True:
        print("no dot!")
        return "无效的开始日期！"

    if (end_year.isdigit() and end_month.isdigit() and end_day.isdigit()) == False:
        print("end date is not digit!")
        return "无效的结束日期！"

    elif '.' in end_year or '.' in end_month or '.' in end_day == True:
        print("no dot!")
        return "无效的结束日期！"

    elif int(begin_year) <= 0 or int(begin_year) > 9999:
        print("begin_year <= 0 or > 9999")
        return "无效的开始年份！"
    elif int(begin_month) <= 0 or int(begin_month) > 12:
        print("begin_month <= 0 or > 12")
        return "无效的开始月份！"
    elif int(begin_day) <= 0 or int(begin_day) > 31:
        print("begin_day <= 0 or > 31")
        return "无效的开始日！"
    elif int(end_year) <= 0 or int(end_year) > 9999:
        print("end_year <= 0 or > 9999")
        return "无效的结束年份！"
    elif int(end_month) <= 0 or int(end_month) > 12:
        print("end_month <= 0 or > 12")
        return "无效的结束月份！"
    elif int(end_day) <= 0 or int(end_day) > 31:
        print("end_day <= 0 or > 31")
        return "无效的结束日！"

    begin = str(int(begin_year)).zfill(4) + str(int(begin_month)).zfill(2) + str(int(begin_day)).zfill(2)
    end = str(int(end_year)).zfill(4) + str(int(end_month)).zfill(2) + str(int(end_day)).zfill(2)

    if isTrueDate(begin) == False:
        print("invalid begin date")
        return "无效的开始日期！"
    if isTrueDate(end) == False:
        print("invalid end date")
        return "无效的结束日期！"
    if begin > end:
        print("begin不能大于end")
        return "无效的日期！"

    begindate = employee_ID + begin
    enddate = employee_ID + end

    con = pymysql.Connect(
        host='XX.XXX.XXX.XX',
        port=XXXX,
        user='XXX',
        passwd='XXXXXXXXXX',
        db='XXX'
    )
    cur = con.cursor()
    sql = 'select work_hour from timecards ' +\
          'where ' + 'employee_ID_date <= ' + "'" + enddate + "' " +\
          "and employee_ID_date >= '" + begindate + "'" +\
          "and charge_number = '" + charge_number + "'"

    print(sql)
    cur.execute(sql)
    value = cur.fetchall()
    print(value)
    cur.close()
    con.close()

    total = 0.0
    for i in value:
        # print(i)
        total += i[0]
    return total


def create_ONE_report(employee_ID, charge_number, workhour, begin_year, begin_month, begin_day, end_year, end_month, end_day):
    if begin_year.isspace() or begin_month.isspace() \
            or begin_day.isspace() or end_year.isspace() \
            or end_month.isspace() or end_day.isspace() == True:
        print('无效输入(1)空格')
        return '输入内容不能全为空格！'
    elif len(begin_year) == 0 or len(begin_month) == 0 \
            or len(begin_day) == 0 or len(end_year) == 0\
            or len(end_month) == 0 or len(end_day) == 0:
        print('无效输入(2)空值')
        return '输入内容不能为空！'
    elif len(begin_year) > 50:
        print('begin_year太长')
        return '开始年份过长！'
    elif len(begin_month) > 50:
        print('begin_month太长')
        return '开始月份过长！'
    elif len(begin_day) > 50:
        print('begin_day太长')
        return '开始日过长！'
    elif len(end_year) > 50:
        print('end_year太长')
        return '结束年份过长！'
    elif len(end_month) > 50:
        print('end_month太长')
        return '结束月份过长！'
    elif len(end_day) > 50:
        print('end_day太长')
        return '结束日过长！'

    if (begin_year.isdigit() and begin_month.isdigit() and begin_day.isdigit()) == False:
        print("begin date is not digit!")
        return "无效的开始日期！"

    elif '.' in begin_year or '.' in begin_month or '.' in begin_day == True:
        print("no dot!")
        return "无效的开始日期！"

    if (end_year.isdigit() and end_month.isdigit() and end_day.isdigit()) == False:
        print("end date is not digit!")
        return "无效的结束日期！"

    elif '.' in end_year or '.' in end_month or '.' in end_day == True:
        print("no dot!")
        return "无效的结束日期！"

    elif int(begin_year) <= 0 or int(begin_year) > 9999:
        print("begin_year <= 0 or > 9999")
        return "无效的开始年份！"
    elif int(begin_month) <= 0 or int(begin_month) > 12:
        print("begin_month <= 0 or > 12")
        return "无效的开始月份！"
    elif int(begin_day) <= 0 or int(begin_day) > 31:
        print("begin_day <= 0 or > 31")
        return "无效的开始日！"
    elif int(end_year) <= 0 or int(end_year) > 9999:
        print("end_year <= 0 or > 9999")
        return "无效的结束年份！"
    elif int(end_month) <= 0 or int(end_month) > 12:
        print("end_month <= 0 or > 12")
        return "无效的结束月份！"
    elif int(end_day) <= 0 or int(end_day) > 31:
        print("end_day <= 0 or > 31")
        return "无效的结束日！"

    begin = str(int(begin_year)).zfill(4) + str(int(begin_month)).zfill(2) + str(int(begin_day)).zfill(2)
    end = str(int(end_year)).zfill(4) + str(int(end_month)).zfill(2) + str(int(end_day)).zfill(2)

    if isTrueDate(begin) == False:
        print("invalid begin date")
        return "无效的开始日期！"
    if isTrueDate(end) == False:
        print("invalid end date")
        return "无效的结束日期！"

    if begin > end:
        print("begin不能大于end")
        return "无效的日期！"

    begindate = employee_ID + begin
    enddate = employee_ID + end
    con = pymysql.Connect(
        host='XX.XXX.XXX.XX',
        port=XXXX,
        user='XXX',
        passwd='XXXXXXXXXX',
        db='XXX'
    )
    cur = con.cursor()
    sql = 'select name from employee_info ' + \
          'where ' + 'employee_ID = ' + "'" + employee_ID + "' "

    print(sql)
    cur.execute(sql)
    value = cur.fetchall()
    print(value[0][0])
    cur.close()
    con.close()
    returnvalue = (value[0][0],)
    returnvalue += (employee_ID,)
    returnvalue += (charge_number,)
    returnvalue += (begin,)
    returnvalue += (end,)
    returnvalue += (workhour,)
    return (returnvalue,)


def create_doc(user, location, title, info_tuple):

    doc = Document()

    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 标题
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    font = run.font

    font.size = Pt(24)
    print("len(info_tuple) = "+str(len(info_tuple)))
    table = doc.add_table(rows=len(info_tuple)+1, cols=6, style="Table Grid")

    # 水平居中
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "员工号"
    table.cell(0, 2).text = "项目编号"
    table.cell(0, 3).text = "开始日期(年月日)"
    table.cell(0, 4).text = "结束日期(年月日)"
    table.cell(0, 5).text = "总工时(时)"

    for i in range(len(info_tuple)):
        table.cell(1 + i, 0).text = info_tuple[i][0]
        table.cell(1 + i, 1).text = info_tuple[i][1]
        table.cell(1 + i, 2).text = info_tuple[i][2]
        table.cell(1 + i, 3).text = info_tuple[i][3]
        table.cell(1 + i, 4).text = info_tuple[i][4]
        table.cell(1 + i, 5).text = str(info_tuple[i][5])

    paragraph = doc.add_paragraph()
    today = datetime.today()
    createtime = str(today.year).zfill(4) + \
               str(today.month).zfill(2) + \
               str(today.day).zfill(2) + \
               str(today.hour).zfill(2) + \
               str(today.minute).zfill(2) + \
               str(today.second).zfill(2)
    run = paragraph.add_run("FROM " + user + "@ACME, " + createtime)
    font = run.font
    font.size = Pt(8)


    # 水平居中
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(location)

    return "YES"


def create_ONE_report_YTD(employee_ID, money, begin_year, begin_month, begin_day, end_year, end_month, end_day):
    if begin_year.isspace() or begin_month.isspace() \
            or begin_day.isspace() or end_year.isspace() \
            or end_month.isspace() or end_day.isspace() == True:
        print('无效输入(1)空格')
        return '输入内容不能全为空格！'
    elif len(begin_year) == 0 or len(begin_month) == 0 \
            or len(begin_day) == 0 or len(end_year) == 0\
            or len(end_month) == 0 or len(end_day) == 0:
        print('无效输入(2)空值')
        return '输入内容不能为空！'
    elif len(begin_year) > 50:
        print('begin_year太长')
        return '开始年份过长！'
    elif len(begin_month) > 50:
        print('begin_month太长')
        return '开始月份过长！'
    elif len(begin_day) > 50:
        print('begin_day太长')
        return '开始日过长！'
    elif len(end_year) > 50:
        print('end_year太长')
        return '结束年份过长！'
    elif len(end_month) > 50:
        print('end_month太长')
        return '结束月份过长！'
    elif len(end_day) > 50:
        print('end_day太长')
        return '结束日过长！'

    if (begin_year.isdigit() and begin_month.isdigit() and begin_day.isdigit()) == False:
        print("begin date is not digit!")
        return "无效的开始日期！"

    elif '.' in begin_year or '.' in begin_month or '.' in begin_day == True:
        print("no dot!")
        return "无效的开始日期！"

    if (end_year.isdigit() and end_month.isdigit() and end_day.isdigit()) == False:
        print("end date is not digit!")
        return "无效的结束日期！"

    elif '.' in end_year or '.' in end_month or '.' in end_day == True:
        print("no dot!")
        return "无效的结束日期！"

    elif int(begin_year) <= 0 or int(begin_year) > 9999:
        print("begin_year <= 0 or > 9999")
        return "无效的开始年份！"
    elif int(begin_month) <= 0 or int(begin_month) > 12:
        print("begin_month <= 0 or > 12")
        return "无效的开始月份！"
    elif int(begin_day) <= 0 or int(begin_day) > 31:
        print("begin_day <= 0 or > 31")
        return "无效的开始日！"
    elif int(end_year) <= 0 or int(end_year) > 9999:
        print("end_year <= 0 or > 9999")
        return "无效的结束年份！"
    elif int(end_month) <= 0 or int(end_month) > 12:
        print("end_month <= 0 or > 12")
        return "无效的结束月份！"
    elif int(end_day) <= 0 or int(end_day) > 31:
        print("end_day <= 0 or > 31")
        return "无效的结束日！"

    begin = str(int(begin_year)).zfill(4) + str(int(begin_month)).zfill(2) + str(int(begin_day)).zfill(2)
    end = str(int(end_year)).zfill(4) + str(int(end_month)).zfill(2) + str(int(end_day)).zfill(2)

    if isTrueDate(begin) == False:
        print("invalid begin date")
        return "无效的开始日期！"
    if isTrueDate(end) == False:
        print("invalid end date")
        return "无效的结束日期！"

    if begin > end:
        print("begin不能大于end")
        return "无效的日期！"

    begindate = employee_ID + begin
    enddate = employee_ID + end
    con = pymysql.Connect(
        host='XX.XXX.XXX.XX',
        port=XXXX,
        user='XXX',
        passwd='XXXXXXXXXX',
        db='XXX'
    )
    cur = con.cursor()
    sql = 'select name from employee_info ' + \
          'where ' + 'employee_ID = ' + "'" + employee_ID + "' "

    print(sql)
    cur.execute(sql)
    value = cur.fetchall()
    print(value[0][0])
    cur.close()
    con.close()
    returnvalue = (value[0][0],)
    returnvalue += (employee_ID,)
    returnvalue += (begin,)
    returnvalue += (end,)
    returnvalue += (money,)
    return (returnvalue,)


def return_pay_YTD(employee_ID, begin_year, begin_month, begin_day, end_year, end_month, end_day):

    con = pymysql.Connect(
        host='XX.XXX.XXX.XX',
        port=XXXX,
        user='XXX',
        passwd='XXXXXXXXXX',
        db='XXX'
    )
    cur = con.cursor()
    sql = "select * from employee_info where employee_type <> 'PA' and "+'employee_ID = ' + "'" + employee_ID + "'"
    print(sql)
    cur.execute(sql)
    value = cur.fetchall()
    cur.close()
    con.close()

    if value == ():
        print('employee_ID不存在')
        return '员工不存在！'

    if employee_ID.isspace() or begin_year.isspace() or begin_month.isspace() \
            or begin_day.isspace() or end_year.isspace() \
            or end_month.isspace() or end_day.isspace() == True:
        print('无效输入(1)空格')
        return '输入内容不能全为空格！'
    elif len(employee_ID) == 0 or len(begin_year) == 0 or len(begin_month) == 0 \
            or len(begin_day) == 0 or len(end_year) == 0\
            or len(end_month) == 0 or len(end_day) == 0:
        print('无效输入(2)空值')
        return '输入内容不能为空！'

    elif len(begin_year) > 50:
        print('begin_year太长')
        return '开始年份过长！'
    elif len(begin_month) > 50:
        print('begin_month太长')
        return '开始月份过长！'
    elif len(begin_day) > 50:
        print('begin_day太长')
        return '开始日过长！'
    elif len(end_year) > 50:
        print('end_year太长')
        return '结束年份过长！'
    elif len(end_month) > 50:
        print('end_month太长')
        return '结束月份过长！'
    elif len(end_day) > 50:
        print('end_day太长')
        return '结束日过长！'

    if (begin_year.isdigit() and begin_month.isdigit() and begin_day.isdigit()) == False:
        print("begin date is not digit!")
        return "无效的开始日期！"

    elif '.' in begin_year or '.' in begin_month or '.' in begin_day == True:
        print("no dot!")
        return "无效的开始日期！"

    if (end_year.isdigit() and end_month.isdigit() and end_day.isdigit()) == False:
        print("end date is not digit!")
        return "无效的结束日期！"

    elif '.' in end_year or '.' in end_month or '.' in end_day == True:
        print("no dot!")
        return "无效的结束日期！"

    elif int(begin_year) <= 0 or int(begin_year) > 9999:
        print("begin_year <= 0 or > 9999")
        return "无效的开始年份！"
    elif int(begin_month) <= 0 or int(begin_month) > 12:
        print("begin_month <= 0 or > 12")
        return "无效的开始月份！"
    elif int(begin_day) <= 0 or int(begin_day) > 31:
        print("begin_day <= 0 or > 31")
        return "无效的开始日！"
    elif int(end_year) <= 0 or int(end_year) > 9999:
        print("end_year <= 0 or > 9999")
        return "无效的结束年份！"
    elif int(end_month) <= 0 or int(end_month) > 12:
        print("end_month <= 0 or > 12")
        return "无效的结束月份！"
    elif int(end_day) <= 0 or int(end_day) > 31:
        print("end_day <= 0 or > 31")
        return "无效的结束日！"

    begin = str(int(begin_year)).zfill(4) + str(int(begin_month)).zfill(2) + str(int(begin_day)).zfill(2)
    end = str(int(end_year)).zfill(4) + str(int(end_month)).zfill(2) + str(int(end_day)).zfill(2)

    if isTrueDate(begin) == False:
        print("invalid begin date")
        return "无效的开始日期！"
    if isTrueDate(end) == False:
        print("invalid end date")
        return "无效的结束日期！"

    if begin > end:
        print("begin不能大于end")
        return "无效的日期！"

    begindate = begin
    enddate = end

    con = pymysql.Connect(
        host='XX.XXX.XXX.XX',
        port=XXXX,
        user='XXX',
        passwd='XXXXXXXXXX',
        db='XXX'
    )
    cur = con.cursor()
    sql = 'select getpaid from payroll where ' + 'enddate <= ' + enddate + ' and enddate >= ' + begindate + " and " +\
          " employee_ID_begin_end like " + "'" + employee_ID + "%' "
    print(sql)
    cur.execute(sql)
    value = cur.fetchall()
    cur.close()
    con.close()
    print(value)

    returnvalue = 0.00
    for i in value:
        print(i)
        returnvalue += i[0]
    print(returnvalue)
    returnvalue = cut(returnvalue, 2)
    return returnvalue


def create_doc_YTD(user, location, title, info_tuple):

    """
        2021/10/12 更新
        功能增强：按 pay period 输出
    """
    new_info_tuple = ()

    for i in range(len(info_tuple)):

        print("info_tuple")
        print(info_tuple[i])

        #new_info_tuple += (info_tuple[i],)

        con = pymysql.Connect(
            host='XX.XXX.XXX.XX',
            port=XXXX,
            user='XXX',
            passwd='XXXXXXXXXX',
            db='XXX'
        )
        cur = con.cursor()
        sql = 'select employee_ID_begin_end, salary, commission, the401k, medical, tax, getpaid from payroll where ' + 'enddate <= ' + info_tuple[i][3] + ' and enddate >= ' + info_tuple[i][2] + " and " + \
              " employee_ID_begin_end like " + "'" + info_tuple[i][1] + "%' "
        print(sql)
        cur.execute(sql)
        value = cur.fetchall()
        cur.close()
        con.close()
        print(value)


        sumsalary = 0.0
        sumcommission = 0.0
        sum401 = 0.0
        summedical = 0.0
        sumtax = 0.0

        for j in value:
            temp = ()
            temp += ("",)
            temp += ("",)
            temp += (j[0][6:14],)
            temp += (j[0][14:22],)
            temp += (str(j[1]),)
            sumsalary += j[1]
            temp += (str(j[2]),)
            sumcommission += j[2]
            temp += (str(j[3]),)
            sum401 += j[3]
            temp += (str(j[4]),)
            summedical += j[4]
            temp += (str(j[5]),)
            sumtax += j[5]
            temp += (str(j[6]),)

            new_info_tuple += (temp,)

        tempname = ()
        tempname += (info_tuple[i][0],)
        # print("!!!!"+info_tuple[i][0])
        tempname += (info_tuple[i][1],)
        # print(tempname)
        tempname += (info_tuple[i][2],)
        tempname += (info_tuple[i][3],)
        tempname += (str(sumsalary),)
        tempname += (str(sumcommission),)
        tempname += (str(sum401),)
        tempname += (str(summedical),)
        tempname += (str(sumtax),)
        tempname += (info_tuple[i][4],)

        new_info_tuple += (tempname,)
        if i != len(info_tuple)-1:
            new_info_tuple += (("-","-","-","-","-","-","-","-","-","-"),)

    print(new_info_tuple)
    info_tuple = new_info_tuple

    """
        以下部分是旧版本代码
    """
    doc = Document()

    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 标题
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    font = run.font

    font.size = Pt(24)
    #print("len(info_tuple) = "+str(len(info_tuple)))
    table = doc.add_table(rows=len(info_tuple)+1, cols=10, style="Table Grid")

    # 水平居中
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table.cell(0, 0).text = '姓名'  # "name"
    table.cell(0, 1).text = '员工号' # "employee ID"
    table.cell(0, 2).text = '开始日期(年月日)'  # "from"
    table.cell(0, 3).text = '结束日期(年月日)'  # "to"
    table.cell(0, 4).text = '底薪(元)'  # "salary"
    table.cell(0, 6).text = '分成(元)'  # "commission"
    table.cell(0, 7).text = '养老金(元)'  # "401k"
    table.cell(0, 8).text = '医保(元)'  # "medical"
    table.cell(0, 5).text = '纳税(元)'  # "tax"
    table.cell(0, 9).text = '到账(元)'  # "payment"

    for i in range(len(info_tuple)):
        table.cell(1 + i, 0).text = info_tuple[i][0]
        table.cell(1 + i, 1).text = info_tuple[i][1]
        table.cell(1 + i, 2).text = info_tuple[i][2]
        table.cell(1 + i, 3).text = info_tuple[i][3]

        table.cell(1 + i, 4).text = info_tuple[i][4]
        table.cell(1 + i, 5).text = info_tuple[i][5]
        table.cell(1 + i, 6).text = info_tuple[i][6]
        table.cell(1 + i, 7).text = str(info_tuple[i][7])
        table.cell(1 + i, 8).text = str(info_tuple[i][8])
        table.cell(1 + i, 9).text = str(info_tuple[i][9])

    paragraph = doc.add_paragraph()
    today = datetime.today()
    createtime = str(today.year).zfill(4) + \
               str(today.month).zfill(2) + \
               str(today.day).zfill(2) + \
               str(today.hour).zfill(2) + \
               str(today.minute).zfill(2) + \
               str(today.second).zfill(2)
    run = paragraph.add_run("FROM " + user + "@ACME, " + createtime)
    font = run.font
    font.size = Pt(8)


    # 水平居中
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(location)

    return "YES"


# print(return_total_hours_worked('000002', '2021', '', '5', '2021', '10', '011'))
# print(return_total_hours_for_a_project('000001', 'C0001', '2021', '8', '5', '2021', '10', '011'))
# print(return_total_hours_for_a_project('000002', 'C0001', '20211003', '20211005'))
# print(create_ONE_report('000001', 'A0001',  '12','1999','01','1', '2020','1','1'))
#a = create_ONE_report('000001', 'A0001', '19990101', '20200101', 12)
#b = create_ONE_report('000001', 'A0001', '19990101', '20200101', 12)
#print(a+b)
#create_doc('000001',"D:\文件生成功能测试.docx",'测试报告', a+a+a+a+a+a+a)

# # print(return_pay_YTD('000003', '2021', '1', '5', '2021', '10', '011'))
# c = create_ONE_report('000003', '10000.00', '10000.00','1999','01','1', '2024','1','1')
# d = create_ONE_report('000002', '10000.00', '10000.00','1999','01','1', '2024','1','1')
# create_doc('000001',"D:\文件生成功能测试.docx",'测试报告', c+d)